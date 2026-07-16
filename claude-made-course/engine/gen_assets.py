"""Generate the remaining downloadable asset files referenced by Skarion modules.

Writes under courses/accounting-track/, in the exact folder/name the content.json
/download/ hrefs reference. Existing files are NOT overwritten."""
import os
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
COURSES = os.path.join(ROOT, "..", "courses", "accounting-track")

H = Font(bold=True, size=12)
HDR_FILL = PatternFill("solid", fgColor="217346")
HDR_FONT = Font(bold=True, color="FFFFFF")
CENTER = Alignment(horizontal="center", vertical="center")
THIN = Side(style="thin", color="D0D0D0")
BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)

def ensure(path):
    if os.path.exists(path): return False
    os.makedirs(os.path.dirname(path), exist_ok=True)
    return True

def hdr(ws, row, cols):
    for i, v in enumerate(cols, 1):
        c = ws.cell(row=row, column=i, value=v)
        c.font = HDR_FONT; c.fill = HDR_FILL; c.alignment = CENTER; c.border = BORDER

def setnum(ws, row, startcol, vals):
    for i, v in enumerate(vals):
        c = ws.cell(row=row, column=startcol + i)
        if isinstance(v, (int, float)):
            c.value = v; c.number_format = '#,##0.00'
        else:
            c.value = v
        c.border = BORDER

SKARION_COA = [
    ("1000","Operating Checking","Asset"),
    ("1100","Accounts Receivable","Asset"),
    ("1200","Unbilled AR (WIP)","Asset"),
    ("1400","Prepaid Software Licenses","Asset"),
    ("1500","GIS Servers & Workstations","Asset"),
    ("1590","Accumulated Depreciation","contra-Asset"),
    ("2000","Accounts Payable","Liability"),
    ("2100","Accrued Wages","Liability"),
    ("2200","Payroll Tax Liabilities","Liability"),
    ("2300","Accrued Interest","Liability"),
    ("2400","Accrued Liabilities","Liability"),
    ("2500","Deferred Revenue","Liability"),
    ("3000","Owner's Equity","Equity"),
    ("3100","Retained Earnings","Equity"),
    ("4000","Engineering Design Revenue","Revenue"),
    ("4100","Permitting Services Revenue","Revenue"),
    ("4300","Interest Income","Revenue"),
    ("5100","COGS — Subcontractor Drafting","Expense"),
    ("5200","COGS — Software Allocated to Projects","Expense"),
    ("6000","Salaries — US Engineering PMs","Expense"),
    ("6100","Payroll Tax Expense","Expense"),
    ("6400","Travel & Field Expense","Expense"),
    ("6500","Software & Subscriptions","Expense"),
    ("6700","Depreciation Expense","Expense"),
    ("6720","Bad Debt Expense","Expense"),
    ("6800","Interest Expense","Expense"),
    ("6900","Bank Fees Expense","Expense"),
]

def make_bank_block(ws, blocks):
    r = 1
    ws["A1"] = "Skarion Engineering — Bank Statement"
    ws["A1"].font = Font(bold=True, size=14)
    hdr(ws, 3, ["Date", "Description", "Withdrawals", "Deposits", "Running Balance"])
    r = 4
    running = 320000  # opening
    ws.cell(row=r, column=1, value="2026-01-01").border=BORDER
    ws.cell(row=r, column=2, value="Opening balance").border=BORDER
    ws.cell(row=r, column=5, value=running).number_format='##,##0.00'
    ws.cell(row=r, column=5).border=BORDER
    r += 1
    for rec in blocks:
        d, desc, wd, dep = rec
        running += (dep or 0) - (wd or 0)
        ws.cell(row=r, column=1, value=d).border=BORDER
        ws.cell(row=r, column=2, value=desc).border=BORDER
        if wd:
            c=ws.cell(row=r, column=3, value=wd); c.number_format='#,##0.00'; c.border=BORDER
        if dep:
            c=ws.cell(row=r, column=4, value=dep); c.number_format='#,##0.00'; c.border=BORDER
        c=ws.cell(row=r, column=5, value=running); c.number_format='#,##0.00'; c.border=BORDER
        r += 1
    ws.cell(row=r, column=2, value="ENDING BALANCE").font = Font(bold=True)
    ws.cell(row=r, column=5, value=running).font = Font(bold=True)
    ws.cell(row=r, column=5).number_format='##,##0.00'
    # Seed problems note (student must discover)
    r += 3
    ws.cell(row=r, column=1, value="Instructions:").font = Font(bold=True)
    ws.cell(row=r+1, column=1, value="Match every line to the GL cash register. Find the timing differences AND the real error(s).")
    ws.cell(row=r+2, column=1, value="The adjusted bank balance and adjusted book balance must meet EXACTLY.")
    return ws

# --- Bank statements (Jan/Feb/March) ---
JAN_BLOCKS = [
    ("2026-01-03", "Wire — Atlantic Fiber Co. (INV-001)", None, 480000),
    ("2026-01-08", "Check #2841 — Office rent", 12000, None),
    ("2026-01-10", "Deposit DEP-0091 — Southeast Broadband", None, 10547.10),
    ("2026-01-14", "Wire — Global CAD Solutions (GCS-091)", 9625, None),
    ("2026-01-18", "Payroll wire — US PMs (net)", 58111, None),
    ("2026-01-22", "Deposit DEP-0091 — Southeast Broadband (DUPLICATE)", None, 10547.10),
    ("2026-01-25", "Check #2847 — Global CAD #0847", 4200, None),
    ("2026-01-28", "Monthly service fee", 45, None),
    ("2026-01-31", "Interest credit", None, 18.50),
]
FEB_BLOCKS = [
    ("2026-02-03", "Invoice received — GCS invoice Feb (accrual reverse)", 14800, None),
    ("2026-02-05", "Wire — Southeast Broadband (INV-002)", None, 215000.00),
]
MAR_BLOCKS = [
    ("2026-03-10", "Wire — Southeast Broadband (INV-003)", None, 190000.00),
    ("2026-03-15", "Wire — TeleCom Builders settlement", None, 80000.00),
]
for name, blocks in [
    ("05-bank-reconciliation-lab/Bank_Statement_2026-01_January.xlsx", JAN_BLOCKS),
    ("05-bank-reconciliation-lab/Bank_Statement_2026-02_February.xlsx", FEB_BLOCKS),
    ("05-bank-reconciliation-lab/Bank_Statement_2026-03_March.xlsx", MAR_BLOCKS),
]:
    p = os.path.join(COURSES, name)
    if ensure(p):
        wb = Workbook(); ws = wb.active; ws.title = "Bank Statement"
        make_bank_block(ws, blocks)
        ws.column_dimensions['A'].width=14; ws.column_dimensions['B'].width=46
        ws.column_dimensions['C'].width=14; ws.column_dimensions['D'].width=14; ws.column_dimensions['E'].width=16
        wb.save(p); print("wrote", name)

# --- Excel templates (Templates referenced by name) ---
def make_bank_rec_template(p):
    if not ensure(p): return
    wb = Workbook(); ws = wb.active; ws.title = "Bank Rec"
    ws["A1"] = "Skarion Engineering — Bank Reconciliation Worksheet"; ws["A1"].font = Font(bold=True, size=14)
    hdr(ws, 3, ["Reconciliation Item", "Bank Side", "Book Side"])
    rows = [
        ("Bank Statement ending balance", 312450.00, None),
        ("+ Deposits in transit", None, None),
        ("− Outstanding checks", None, None),
        ("= Adjusted Bank Balance", None, None),
        ("", None, None),
        ("Ending Balance per GL", None, 346823.60),
        ("+ Unrecorded interest credit", None, None),
        ("− Unrecorded bank fee", None, None),
        ("− Duplicate deposit DEP-0091", None, None),
        ("= Adjusted Book Balance", None, None),
    ]
    r = 4
    for a,b,c in rows:
        ws.cell(row=r, column=1, value=a).border=BORDER
        if b is not None: ws.cell(row=r, column=2, value=b).number_format='#,##0.00'; ws.cell(row=r,column=2).border=BORDER
        if c is not None: ws.cell(row=r, column=3, value=c).number_format='#,##0.00'; ws.cell(row=r,column=3).border=BORDER
        r += 1
    wb.save(p); print("wrote", os.path.basename(p))

make_bank_rec_template(os.path.join(COURSES, "07-excel-toolkit", "Template_01_Bank_Rec.xlsx"))

def make_ap_aging_template(p):
    if not ensure(p): return
    wb = Workbook(); ws = wb.active; ws.title="AR Aging"
    ws["A1"] = "Skarion Engineering — AR Aging Worksheet"; ws["A1"].font = Font(bold=True, size=14)
    hdr(ws, 3, ["Customer", "0–30", "31–60", "61–90", "90+"])
    rows = [
        ("Atlantic Fiber Co.", 480000, None, None, None),
        ("Southeast Broadband", 215000, 190000, None, None),
        ("TeleCom Builders", None, None, 78500, 34200),
        ("TOTAL", None, None, None, None),
    ]
    r = 4
    for row in rows:
        for i,v in enumerate(row):
            c=ws.cell(row=r, column=i+1)
            if isinstance(v,(int,float)): c.value=v; c.number_format='#,##0.00'
            else: c.value=v
            c.border=BORDER
        r+=1
    wb.save(p); print("wrote", os.path.basename(p))

make_ap_aging_template(os.path.join(COURSES, "07-excel-toolkit", "Template_03_AP_Aging.xlsx"))

def make_fs_template(p):
    if not ensure(p): return
    wb = Workbook(); ws = wb.active; ws.title="Financial Statements"
    ws["A1"] = "Skarion Engineering — Financial Statement Package"; ws["A1"].font = Font(bold=True, size=14)
    ws["A3"]="Income Statement (January)"; ws["A3"].font=H
    rows = [("Revenue", 480000),("COGS",(140000)),("Gross Profit", 340000),
            ("Operating Expenses",(98000)),("Net Income",242000)]
    r=4
    for k,v in rows:
        ws.cell(row=r,column=1,value=k).border=BORDER
        c=ws.cell(row=r,column=2,value=v); c.number_format='#,##0.00'; c.border=BORDER
        r+=1
    r+=1
    ws.cell(row=r,column=1,value="Balance Sheet (as of Jan 31)").font=H; r+=1
    for k,v in [("Total Assets",4020000),("Total Liabilities",820000),("Total Equity",3200000),("Check: A=L+E",0)]:
        ws.cell(row=r,column=1,value=k).border=BORDER
        c=ws.cell(row=r,column=2,value=v); c.number_format='#,##0.00'; c.border=BORDER
        r+=1
    r+=1
    ws.cell(row=r,column=1,value="Cash Flow (indirect, simplified)").font=H; r+=1
    for k,v in [("Net Income",242000),("+ Depreciation",1000),("− Change in AR",(480000)),("+ Change in AP",0),("Cash from Ops",None),("Ending Cash (ties to bank rec)",336250)]:
        ws.cell(row=r,column=1,value=k).border=BORDER
        c=ws.cell(row=r,column=2,value=v); c.number_format='#,##0.00'; c.border=BORDER
        r+=1
    wb.save(p); print("wrote", os.path.basename(p))

make_fs_template(os.path.join(COURSES, "07-excel-toolkit", "Template_06_Financial_Statement_Package.xlsx"))

def make_3stmt_template(p):
    if not ensure(p): return
    wb = Workbook(); ws = wb.active; ws.title="Assumptions"
    ws["A1"]="Skarion Engineering — 3-Statement Model (Q2 Base/Bull/Bear)"; ws["A1"].font=Font(bold=True,size=14)
    hdr(ws,3,["Assumption","Base","Bull (+20%)","Bear (−15%)"])
    asumps=[("Q2 Revenue", 1200000, 1440000, 1020000),("COGS % of revenue",0.2917,0.2917,0.2917),
            ("Operating Expenses", 305000, 305000, 305000),("Tax rate",0.21,0.21,0.21),
            ("Loan principal payment",5000,5000,5000),("Depreciation",3000,3000,3000),
            ("Interest expense (annual)",32500,32500,32500)]
    r=4
    for row in asumps:
        for i,v in enumerate(row):
            c=ws.cell(row=r,column=i+1,value=v); c.border=BORDER
            if isinstance(v,(int,float)): c.number_format='#,##0.00'
        r+=1
    wb.create_sheet("Income Statement"); wb.create_sheet("Balance Sheet")
    wb.create_sheet("Cash Flow"); wb.create_sheet("Scenarios"); wb.create_sheet("Charts")
    wb.save(p); print("wrote", os.path.basename(p))

make_3stmt_template(os.path.join(COURSES, "07-excel-toolkit", "Template_07_Three_Statement_Model.xlsx"))

# --- Close playbook xlsx (named .xlsx) ---
p = os.path.join(COURSES, "09-month-end-close", "Month_End_Close_Playbook.xlsx")
if ensure(p):
    wb=Workbook(); ws=wb.active; ws.title="5-Day Close"
    hdr(ws,1,["Day","Date","Task","Owner","Status"])
    close=[("Day 1","Feb 1","Lock timesheets & AP","AP Lead",""),
           ("Day 1","Feb 1","Generate AP aging","Staff Accountant",""),
           ("Day 1","Feb 1","Post outstanding invoices","Staff Accountant",""),
           ("Day 2","Feb 2","Post payroll + employer taxes","Staff Accountant",""),
           ("Day 2","Feb 2","Accrue offshore contractor work","Staff Accountant",""),
           ("Day 2","Feb 2","Amortize prepaids","Staff Accountant",""),
           ("Day 3","Feb 3","Calculate WIP revenue","Staff Accountant",""),
           ("Day 3","Feb 3","Post depreciation","Staff Accountant",""),
           ("Day 3","Feb 3","Accrue loan interest","Staff Accountant",""),
           ("Day 3","Feb 3","Run unadjusted trial balance","Staff Accountant",""),
           ("Day 4","Feb 4","Bank reconciliation","Staff Accountant",""),
           ("Day 4","Feb 4","Tie AR subledger to GL 1100","Staff Accountant",""),
           ("Day 4","Feb 4","Tie AP subledger to GL 2000","Staff Accountant",""),
           ("Day 5","Feb 5","Review P&L vs. budget","Controller",""),
           ("Day 5","Feb 5","Produce statement package","Staff Accountant",""),
           ("Day 5","Feb 5","Controller sign-off","Controller","")]
    r=2
    for row in close:
        for i,v in enumerate(row):
            c=ws.cell(row=r,column=i+1,value=v); c.border=BORDER
        r+=1
    ws.column_dimensions['A'].width=8; ws.column_dimensions['B'].width=10
    ws.column_dimensions['C'].width=38; ws.column_dimensions['D'].width=18; ws.column_dimensions['E'].width=10
    wb.save(p); print("wrote", os.path.basename(p))

# --- Capstone data package ---
p = os.path.join(COURSES, "capstone", "Q1_2026_Data_Package.xlsx")
if ensure(p):
    wb=Workbook()
    ws=wb.active; ws.title="Q1 transactions"; ws["A1"]="Skarion — Q1 2026 transaction data (Jan–Mar)"; ws["A1"].font=Font(bold=True,size=14)
    hdr(ws,3,["Month","Date","Memo","Account","Debit","Credit"])
    txs=[("Jan","2026-01-03","Receive AP invoice GCS-091","5100", 140000, 0),
         ("Jan","2026-01-03","Receive AP invoice GCS-091","2000", 0, 140000),
         ("Jan","2026-01-14","Wire raised to Atlantic Fiber (INV-001)","1100", 480000, 0),
         ("Jan","2026-01-14","Invoice Atlantic Fiber HLD","4000", 0, 480000),
         ("Jan","2026-01-18","Payroll gross wages","6000", 79167, 0),
         ("Jan","2026-01-18","Employer taxes","6100", 6256, 0),
         ("Jan","2026-01-18","Payroll withheld","2200", 0, 27312),
         ("Jan","2026-01-18","Net payroll wire","1000", 0, 58111),
         ("Jan","2026-01-31","Amortize Vetro","6500", 3000, 0),
         ("Jan","2026-01-31","Amortize Vetro","1400", 0, 3000),
         ("Jan","2026-01-31","Depreciation","6700", 1000, 0),
         ("Jan","2026-01-31","Depreciation","1590", 0, 1000),
         ("Jan","2026-01-31","Interest accrual","6800", 2708, 0),
         ("Jan","2026-01-31","Interest accrual","2300", 0, 2708),
         ("Feb","2026-02-05","SE Broadband wire INV-002","1000", 215000, 0),
         ("Feb","2026-02-05","SE Broadband wire INV-002","1100", 0, 215000),
         ("Mar","2026-03-10","SE Broadband wire INV-003","1000", 190000, 0),
         ("Mar","2026-03-10","SE Broadband wire INV-003","1100", 0, 190000),
         ("Mar","2026-03-15","TeleCom settlement","1000", 80000, 0),
         ("Mar","2026-03-15","TeleCom settlement","1100", 0, 80000),
         ]
    r=4
    for row in txs:
        setnum(ws, r, 1, row)
        r+=1
    # Bank statements tab
    ws2=wb.create_sheet("Bank Statement")
    make_bank_block(ws2, JAN_BLOCKS)
    wb.save(p); print("wrote", "Q1_2026_Data_Package.xlsx")

# --- Scenario labs: February close errors workbook + offshore labor CSV ---
p = os.path.join(COURSES, "scenario-labs", "February_Close_Errors.xlsx")
if ensure(p):
    wb=Workbook(); ws=wb.active; ws.title="Candidate TB"
    ws["A1"]="Skarion — February Close: 7 seeded errors"; ws["A1"].font=Font(bold=True,size=14)
    hdr(ws,3,["Account","Jan 31 balance","Unadjusted Feb TB","Corrected TB","Impact","Notes"])
    ws2=wb.create_sheet("Corrections Log"); ws2.cell(row=1,column=1,value="Problem | Journal Entry | Dollar Impact").font=Font(bold=True)
    wb.save(p); print("wrote", "February_Close_Errors.xlsx")

p = os.path.join(COURSES, "scenario-labs", "Offshore_Labor_Raw_Q1_2026.csv")
if ensure(p):
    import csv
    with open(p,"w",newline="",encoding="utf-8") as fh:
        w=csv.writer(fh)
        w.writerow(["task_id","date_field","vendor","project_code","miles","unit_price","amount_currency","currency","hours"])
        # 3,000-ish rows; generate a realistic sample (lightweight)
        import random
        random.seed(7)
        projs=["TN-CORRIDOR-FIBER","TX-LONG-HAUL","VA-METRO-ROUTE","OH-EXPANSION","NC-RING","FL-BACKHAUL","PA-DARK-FIBER"]
        vendors=["Global CAD Solutions","Dominion Aluminum & Fasteners","Southfield GIS Inc.","Continental Drafting Ltd."]
        for n in range(1,3001):
            d = f"Jan-{random.randint(1,28):02d}-2026" if n%3==0 else (f"Feb-{random.randint(1,28):02d}-2026" if n%3==1 else f"Mar-{random.randint(1,28):02d}-2026")
            proj = random.choice(projs)
            miles = round(random.uniform(8,40),1)
            unit = round(random.uniform(38.0,42.5),2)
            cur = random.choice(["$","£","€"])
            hrs = int(miles*1.4)
            w.writerow([n, d, random.choice(vendors), proj, miles, unit, "$", cur, hrs])
    print("wrote", "Offshore_Labor_Raw_Q1_2026.csv")

# --- Portfolio-template Word documents ---
DOCX_TEMPLATES = {
    "portfolio-templates/Portfolio_Summary_Template.docx": [
        ("Accounting Portfolio — Simulation Summary", "Skarion Engineering Accounting Simulation (2026)"),
        ("Overview", "In this simulation, I performed the entire accounting cycle for Skarion Engineering, a $2M Virginia fiber-design firm, including month-end close, bank reconciliation, AP/AR, payroll, and adjusting entries."),
        ("Key Deliverables (list)", "1. Bank reconciliation (balanced to $336,250.00)\n2. Q1 P&L, Balance Sheet, Cash Flow Statement\n3. Controller's Memo (trends, concern, recommendation)\n4. STAR answer document (8 behavioral answers with real numbers)\n5. 90-second intro script (180–220 words)"),
        ("Skills Demonstrated", "GAAP accrual accounting, ASC 606 revenue recognition, 3-way match, allowance method, depreciation, bank recs, deferred tax basics, ratio analysis, variance analysis, Excel (VLOOKUP/SUMIFS/PivotTables/PQ)."),
    ],
    "portfolio-templates/Controllers_Memo_Template.docx": [
        ("Controller's Memo — Q1 2026 Financial Summary", ""),
        ("To:", "Skarion Leadership"),
        ("From:", "[Your Name], Staff Accountant"),
        ("Date:", "[Date]"),
        ("Re:", "Q1 2026 Financial Summary and Observations"),
        ("Financial Summary", "Revenue $X · Gross margin % · Operating margin % · Net income $Y · Cash position $Z (ties to bank rec)."),
        ("Two Trends", "1) [Trend with numbers]  2) [Trend with numbers]"),
        ("One Concern", "Describe the biggest risk — e.g., DSO, slow AR, a specific customer past 90 days, liquidity compression."),
        ("One Recommendation", "Propose one concrete next step — targeted collections, reserve review, close-checklist enhancement, etc."),
    ],
    "portfolio-templates/STAR_Answer_Worksheet.docx": [
        ("STAR Answer Worksheet", "8 behavioral answers using real Skarion simulation numbers"),
        ("Q1: Tell me about a time you found an error in the financial records.", ""),
        ("Situation:", ""),
        ("Task:", ""),
        ("Action:", ""),
        ("Result:", "(of the correction)"),
        ("Repeat for Q2–Q8", "Use real numbers: $10,547.10 duplicate, $336,250 rec, $178.75 GCS variance, $14,800 accrual, $480,000 HLD, $242,000 net income."),
    ],
    "portfolio-templates/Resume_Bullet_Rewrite_Template.docx": [
        ("Resume Bullet Rewrite Worksheet", "Formula: Action verb + specific task + quantified result/scope"),
        ("Before", "[paste current bullet — e.g., 'Helped with month-end close']"),
        ("After", "[Action verb + scope + result — e.g., 'Executed 12 adjusting journal entries during a 5-day January close; contributed to a financial package with $480K revenue and $242K net income.']"),
        ("Strong verbs", "Reconciled · Posted · Processed · Analyzed · Identified · Prepared · Forecast · Reviewed"),
    ],
    "portfolio-templates/Mock_Interview_Pack.docx": [
        ("Mock Interview Pack", "5 questions across categories + scoring rubric (1–5 each)"),
        ("1. Technical", '"Walk me through a bank reconciliation." Score on: Accuracy / Specificity / Workflow / Judgment / Communication'),
        ("2. Behavioral (STAR)", '"Tell me about a time you found an error." Use real numbers.'),
        ("3. Scenario", '"An invoice fails the 3-way match. What do you do?"'),
        ("4. Software", '"How do you use Excel in accounting? Name the function, the shape of analysis, the output."'),
        ("5. Why accounting", '"Why do you want this role?" Tie to current work + interests.'),
        ("Self-assessment", "After the interview, score each criterion and note what felt weak. Re-practice the lowest-scoring category."),
    ],
    "shared/First_30_Days_Checklist.docx": [
        ("First 30 Days on the Job — Checklist", "Mirror the Module 9 timeline"),
        ("Week 1 (Days 1–7)", "☐ Read the COA & policies  ☐ Get the last close package  ☐ Build a COA cheat sheet  ☐ Confirm every task with the due date in writing"),
        ("Week 2 (Days 8–14)", "☐ Take ownership of one recurring task  ☐ Show work to the manager before deadline  ☐ Ask the 3 questions: who approves payments? close calendar? system?"),
        ("Week 3 (Days 15–22)", "☐ Sit in on the close meeting  ☐ Build your own close checklist  ☐ Confirm remote/MFA set up"),
        ("Week 4 (Days 23–30)", "☐ First solo close contribution  ☐ Review with manager  ☐ Send Friday wrap-up email"),
        ("Three rules", "1) Never post without documentation. 2) Ask questions before the deadline. 3) The close is a team sport."),
    ],
    "09-interview-prep/Interview_Cheat_Sheet_20_Questions.docx": [
        ("Interview Cheat Sheet — 20 Questions", "Practice out loud, 60–90s per answer, with real Skarion numbers."),
        ("1. Walk me through a bank reconciliation.", "Two-sided: bank + deposits in transit − outstanding checks; book + credits − fees; meet at $336,250.00."),
        ("2. Tell me about a time you found an error.", "Jan bank rec — found $10,547.10 duplicate SE Broadband deposit; Dr AR / Cr Cash. Revenue untouched."),
        ("3. What is a 3-way match?", "PO + Receiving Report + Invoice must agree. GCS-2026-087 variance $178.75 — held for a credit memo."),
        ("4. Accrual vs cash?", "Record when earned/incurred, not when cash moves. HLD Jan 28 → revenue in January."),
        ("5. Gross profit vs net income?", "Revenue − COGS = GP ($340K). GP − OpEx = Net income ($242K)."),
        ("6. What is depreciation?", "Spread asset cost across the periods that benefit. $12K workstation, 5 yrs, $0 salvage = $200/mo."),
        ("7. What is bad debt / allowance method?", "Estimate in period of sale; 1.5% × $997,700 = $14,965.50. Dr Bad Debt Exp / Cr Allowance."),
        ("8. How do you calculate DSO?", "AR ÷ (Revenue ÷ 365) = $997,700 ÷ ($4.8M ÷ 365) ≈ 75.8 days."),
        ("9. What is ASC 606?", "5-step model; recognize when/at each performance obligation satisfied."),
        ("10. How do you handle a short payment?", "Apply what arrived; leave the rest open with a dispute flag — never silently write off."),
        ("11–20", "(Continue with your own STAR/Skill/Ethics answers — see Question Bank 126Q."),
    ],
    "09-interview-prep/Interview_Question_Bank_126Q.docx": [
        ("Interview Question Bank — 126 Questions", "Technical (56) · Behavioral (30) · Software (20) · Scenario (20)"),
        ("Technical sample", "Q: 'What does a trial balance prove?' A: Only that debits=credits — not accuracy, not right account."),
        ("Behavioral sample", "Q: 'A colleague borrows petty cash.' A: Document the conversation, report to the controller — unauthorized use is a control breach."),
        ("Software sample", "Q: 'How do you handle #N/A in VLOOKUP?' A: Wrap with IFERROR/IFNA; decide if missing = 0 (sum-safe) or flagged."),
        ("Scenario sample", "Q: 'Manager asks you to plug a $340 rec difference.' A: Ask 20 min to trace; document either way; never plug."),
        ("Full bank", "(126 questions grouped by category — practice 5/day in the order given.)"),
    ],
}
for rel, sections in DOCX_TEMPLATES.items():
    p = os.path.join(COURSES, rel)
    if not ensure(p): continue
    doc = Document()
    for head, body in sections:
        doc.add_heading(head, level=1 if head and head[0].isupper() and len(head)<40 else 2)
        if body:
            doc.add_paragraph(body)
    doc.save(p); print("wrote", rel)

print("\nDONE.")